from __future__ import annotations

from io import BytesIO
from flask import Blueprint, jsonify, request, send_file
from sqlalchemy import select
from docx import Document

from ..database import SessionLocal
from ..models import Chapter, Novel, Character, Idea, ChapterVersion


novel_bp = Blueprint("novels", __name__, url_prefix="/api")


@novel_bp.get("/stats")
def get_stats():
    with SessionLocal() as db:
        novel_count = db.query(Novel).count()
        chapter_count = db.query(Chapter).count()
        character_count = db.query(Character).count()
        # 简单估算字数
        all_chapters = db.scalars(select(Chapter)).all()
        word_count = sum(len(c.content) for c in all_chapters if c.content)
        
    return jsonify({
        "code": "OK", 
        "data": {
            "novel_count": novel_count,
            "chapter_count": chapter_count,
            "character_count": character_count,
            "word_count": word_count
        }
    })


@novel_bp.get("/novels")
def list_novels():
    with SessionLocal() as db:
        novels = db.scalars(select(Novel).order_by(Novel.updated_at.desc())).all()
        data = [
            {
                "id": n.id,
                "title": n.title,
                "summary": n.summary,
                "tags": n.tags,
                "updated_at": n.updated_at.isoformat(),
            }
            for n in novels
        ]
    return jsonify({"code": "OK", "data": data})


@novel_bp.post("/novels")
def create_novel():
    body = request.get_json(silent=True) or {}
    title = str(body.get("title", "")).strip()
    if not title:
        return jsonify({"code": "INVALID_INPUT", "message": "标题不能为空"}), 400
    summary = body.get("summary")
    tags = body.get("tags")
    with SessionLocal() as db:
        novel = Novel(owner_id=1, title=title, summary=summary, tags=tags)
        db.add(novel)
        db.commit()
        db.refresh(novel)
    return jsonify({"code": "OK", "data": {"id": novel.id}})


@novel_bp.put("/novels/<int:novel_id>")
def update_novel(novel_id: int):
    body = request.get_json(silent=True) or {}
    title = body.get("title")
    summary = body.get("summary")
    tags = body.get("tags")
    
    with SessionLocal() as db:
        novel = db.get(Novel, novel_id)
        if not novel:
            return jsonify({"code": "NOT_FOUND", "message": "小说不存在"}), 404
            
        if isinstance(title, str) and title.strip():
            novel.title = title.strip()
        if isinstance(summary, str):
            novel.summary = summary
        if isinstance(tags, str):
            novel.tags = tags
            
        db.add(novel)
        db.commit()
    return jsonify({"code": "OK"})


@novel_bp.delete("/novels/<int:novel_id>")
def delete_novel(novel_id: int):
    with SessionLocal() as db:
        novel = db.get(Novel, novel_id)
        if novel:
            # SQLAlchemy cascade delete should handle children if configured, 
            # but usually it's better to rely on DB constraints or manual cleanup if needed.
            # Assuming models are defined with cascade or we just delete the novel and DB handles it (or orphan records remain).
            # For simplicity let's just delete the novel object.
            db.delete(novel)
            db.commit()
    return jsonify({"code": "OK"})


@novel_bp.get("/novels/<int:novel_id>/export")
def export_novel(novel_id: int):
    fmt = request.args.get("format", "txt")
    with SessionLocal() as db:
        novel = db.get(Novel, novel_id)
        if not novel:
            return jsonify({"code": "NOT_FOUND", "message": "小说不存在"}), 404
        
        chapters = db.scalars(
            select(Chapter).where(Chapter.novel_id == novel_id).order_by(Chapter.order_index.asc())
        ).all()

    filename = f"{novel.title}.{fmt}"
    
    if fmt == "txt":
        buffer = BytesIO()
        text_content = f"《{novel.title}》\n\n简介：{novel.summary}\n\n"
        for chapter in chapters:
            text_content += f"\n\n第{chapter.order_index}章 {chapter.title}\n\n{chapter.content}\n"
        buffer.write(text_content.encode("utf-8"))
        buffer.seek(0)
        return send_file(buffer, as_attachment=True, download_name=filename, mimetype="text/plain")

    elif fmt == "docx":
        doc = Document()
        doc.add_heading(novel.title, 0)
        if novel.summary:
            doc.add_paragraph(f"简介：{novel.summary}")
        
        for chapter in chapters:
            doc.add_page_break()
            doc.add_heading(f"第{chapter.order_index}章 {chapter.title}", level=1)
            doc.add_paragraph(chapter.content)
            
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return send_file(
            buffer, 
            as_attachment=True, 
            download_name=filename, 
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    else:
        return jsonify({"code": "INVALID_FORMAT", "message": "不支持的导出格式"}), 400


@novel_bp.get("/novels/<int:novel_id>/ideas")
def list_ideas(novel_id: int):
    with SessionLocal() as db:
        ideas = db.scalars(
            select(Idea).where(Idea.novel_id == novel_id).order_by(Idea.created_at.desc())
        ).all()
        data = [
            {"id": i.id, "content": i.content, "idea_type": i.idea_type, "created_at": i.created_at.isoformat()}
            for i in ideas
        ]
    return jsonify({"code": "OK", "data": data})


@novel_bp.post("/novels/<int:novel_id>/ideas")
def create_idea(novel_id: int):
    body = request.get_json(silent=True) or {}
    content = str(body.get("content", "")).strip()
    idea_type = str(body.get("idea_type", "general"))
    
    if not content:
        return jsonify({"code": "INVALID_INPUT", "message": "内容不能为空"}), 400
        
    with SessionLocal() as db:
        idea = Idea(novel_id=novel_id, content=content, idea_type=idea_type)
        db.add(idea)
        db.commit()
        db.refresh(idea)
    return jsonify({"code": "OK", "data": {"id": idea.id}})


@novel_bp.delete("/ideas/<int:idea_id>")
def delete_idea(idea_id: int):
    with SessionLocal() as db:
        idea = db.get(Idea, idea_id)
        if idea:
            db.delete(idea)
            db.commit()
    return jsonify({"code": "OK"})


@novel_bp.get("/novels/<int:novel_id>/chapters")
def list_chapters(novel_id: int):
    with SessionLocal() as db:
        chapters = db.scalars(
            select(Chapter).where(Chapter.novel_id == novel_id).order_by(Chapter.order_index.asc())
        ).all()
        data = [
            {"id": c.id, "title": c.title, "order_index": c.order_index, "updated_at": c.updated_at.isoformat()}
            for c in chapters
        ]
    return jsonify({"code": "OK", "data": data})


@novel_bp.post("/novels/<int:novel_id>/chapters")
def create_chapter(novel_id: int):
    body = request.get_json(silent=True) or {}
    title = str(body.get("title", "")).strip() or "未命名章节"
    with SessionLocal() as db:
        last = db.scalar(
            select(Chapter).where(Chapter.novel_id == novel_id).order_by(Chapter.order_index.desc())
        )
        next_index = (last.order_index + 1) if last else 1
        chapter = Chapter(novel_id=novel_id, title=title, order_index=next_index, content="")
        db.add(chapter)
        db.commit()
        db.refresh(chapter)
    return jsonify({"code": "OK", "data": {"id": chapter.id}})


@novel_bp.get("/chapters/<int:chapter_id>")
def get_chapter(chapter_id: int):
    with SessionLocal() as db:
        chapter = db.get(Chapter, chapter_id)
        if not chapter:
            return jsonify({"code": "NOT_FOUND", "message": "章节不存在"}), 404
        data = {"id": chapter.id, "novel_id": chapter.novel_id, "title": chapter.title, "content": chapter.content}
    return jsonify({"code": "OK", "data": data})


@novel_bp.put("/chapters/<int:chapter_id>")
def update_chapter(chapter_id: int):
    body = request.get_json(silent=True) or {}
    content = body.get("content")
    title = body.get("title")
    with SessionLocal() as db:
        chapter = db.get(Chapter, chapter_id)
        if not chapter:
            return jsonify({"code": "NOT_FOUND", "message": "章节不存在"}), 404
        if isinstance(content, str):
            chapter.content = content
        if isinstance(title, str) and title.strip():
            chapter.title = title.strip()
        db.add(chapter)
        db.commit()
    return jsonify({"code": "OK"})


@novel_bp.delete("/chapters/<int:chapter_id>")
def delete_chapter(chapter_id: int):
    with SessionLocal() as db:
        chapter = db.get(Chapter, chapter_id)
        if chapter:
            db.delete(chapter)
            db.commit()
    return jsonify({"code": "OK"})


@novel_bp.delete("/characters/<int:char_id>")
def delete_character(char_id: int):
    with SessionLocal() as db:
        char = db.get(Character, char_id)
        if char:
            db.delete(char)
            db.commit()
    return jsonify({"code": "OK"})


@novel_bp.get("/novels/<int:novel_id>/characters")
def list_characters(novel_id: int):
    with SessionLocal() as db:
        chars = db.scalars(
            select(Character).where(Character.novel_id == novel_id).order_by(Character.created_at.desc())
        ).all()
        data = [
            {"id": c.id, "name": c.name, "profile": c.profile}
            for c in chars
        ]
    return jsonify({"code": "OK", "data": data})


@novel_bp.post("/novels/<int:novel_id>/characters")
def create_character(novel_id: int):
    body = request.get_json(silent=True) or {}
    name = str(body.get("name", "")).strip()
    profile = str(body.get("profile", ""))
    
    if not name:
        return jsonify({"code": "INVALID_INPUT", "message": "姓名不能为空"}), 400
        
    with SessionLocal() as db:
        char = Character(novel_id=novel_id, name=name, profile=profile)
        db.add(char)
        db.commit()
        db.refresh(char)
    return jsonify({"code": "OK", "data": {"id": char.id}})


@novel_bp.put("/characters/<int:char_id>")
def update_character(char_id: int):
    body = request.get_json(silent=True) or {}
    name = body.get("name")
    profile = body.get("profile")
    
    with SessionLocal() as db:
        char = db.get(Character, char_id)
        if not char:
             return jsonify({"code": "NOT_FOUND", "message": "角色不存在"}), 404
             
        if isinstance(name, str) and name.strip():
            char.name = name.strip()
        if isinstance(profile, str):
            char.profile = profile
            
        db.add(char)
        db.commit()
    return jsonify({"code": "OK"})


@novel_bp.get("/chapters/<int:chapter_id>/versions")
def list_chapter_versions(chapter_id: int):
    with SessionLocal() as db:
        versions = db.scalars(
            select(ChapterVersion)
            .where(ChapterVersion.chapter_id == chapter_id)
            .order_by(ChapterVersion.created_at.desc())
        ).all()
        data = [
            {
                "id": v.id,
                "content": v.content, # Optionally exclude content for list if too large
                "note": v.note,
                "created_at": v.created_at.isoformat()
            }
            for v in versions
        ]
    return jsonify({"code": "OK", "data": data})


@novel_bp.post("/chapters/<int:chapter_id>/versions")
def create_chapter_version(chapter_id: int):
    body = request.get_json(silent=True) or {}
    note = body.get("note")
    
    with SessionLocal() as db:
        chapter = db.get(Chapter, chapter_id)
        if not chapter:
            return jsonify({"code": "NOT_FOUND", "message": "章节不存在"}), 404
            
        version = ChapterVersion(
            chapter_id=chapter_id,
            content=chapter.content,
            note=note
        )
        db.add(version)
        db.commit()
        db.refresh(version)
        
    return jsonify({"code": "OK", "data": {"id": version.id}})


@novel_bp.post("/chapters/<int:chapter_id>/restore/<int:version_id>")
def restore_chapter_version(chapter_id: int, version_id: int):
    with SessionLocal() as db:
        chapter = db.get(Chapter, chapter_id)
        version = db.get(ChapterVersion, version_id)
        
        if not chapter or not version:
            return jsonify({"code": "NOT_FOUND", "message": "章节或版本不存在"}), 404
            
        if version.chapter_id != chapter_id:
             return jsonify({"code": "INVALID_OPERATION", "message": "版本不属于该章节"}), 400
             
        # Optional: Create a backup of current state before restoring?
        # For now, just overwrite
        chapter.content = version.content
        db.add(chapter)
        db.commit()
        
    return jsonify({"code": "OK"})


@novel_bp.delete("/versions/<int:version_id>")
def delete_chapter_version(version_id: int):
    with SessionLocal() as db:
        version = db.get(ChapterVersion, version_id)
        if version:
            db.delete(version)
            db.commit()
    return jsonify({"code": "OK"})
